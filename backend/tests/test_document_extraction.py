import fitz
from docx import Document

from app.documents.extraction import extract_document_text, normalize_text, text_quality


def test_text_pdf_uses_text_layer(tmp_path):
    path = tmp_path / "resume.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Jane Doe\njane@example.com\nSkills\nPython, FastAPI")
    document.save(path)
    document.close()
    result = extract_document_text(path)
    assert result.pages == 1
    assert result.method == "pdf_text"
    assert "jane@example.com" in result.text
    assert result.quality > 0


def test_malformed_pdf_is_reported(tmp_path):
    path = tmp_path / "broken.pdf"
    path.write_bytes(b"not a pdf")
    result = extract_document_text(path)
    assert result.method == "failed"
    assert result.error
    assert result.quality == 0


def test_docx_and_normalization(tmp_path):
    path = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("Experience   5 years")
    document.save(path)
    result = extract_document_text(path)
    assert result.method == "docx_text"
    assert "Experience 5 years" in result.text
    assert text_quality(normalize_text("  useful   text ")) > 0
